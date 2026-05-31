import re
import io
import base64
import json
import httpx
import logging
import fitz  # PyMuPDF
from docx import Document
from typing import Dict, Any, List
from app.config import settings

logger = logging.getLogger(__name__)

class ResumeParser:
    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract plain text from PDF using PyMuPDF"""
        text = ""
        try:
            pdf_file = io.BytesIO(file_bytes)
            doc = fitz.open(stream=pdf_file, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {e}")
        return text

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract plain text from DOCX using python-docx"""
        text = ""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {e}")
        return text

    def extract_text_via_gemini_ocr(self, file_bytes: bytes) -> str:
        """Use Gemini 2.5 Flash to perform OCR/vision-based text extraction on a scanned PDF"""
        if not settings.GEMINI_API_KEY:
            raise ValueError("Gemini API key is required to parse scanned PDF files, but it is missing.")
            
        encoded_pdf = base64.b64encode(file_bytes).decode("utf-8")
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={settings.GEMINI_API_KEY}"
        headers = {"Content-Type": "application/json"}
        
        prompt = (
            "You are an expert ATS resume parsing engine. Perform high-fidelity OCR on this PDF. "
            "Extract all content, work experience, education, skills, projects, certifications, and contact details. "
            "Format the output as a clean, structured Markdown document. Do not summarize or omit anything. "
            "Maintain chronological ordering and exact text."
        )
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "inlineData": {
                                "mimeType": "application/pdf",
                                "data": encoded_pdf
                            }
                        },
                        {
                            "text": prompt
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1
            }
        }
        
        try:
            with httpx.Client() as client:
                r = client.post(url, json=payload, headers=headers, timeout=45.0)
                if r.status_code == 200:
                    data = r.json()
                    response_text = data["contents"][0]["parts"][0]["text"]
                    return response_text
                else:
                    raise ValueError(f"Gemini OCR API returned code {r.status_code}: {r.text}")
        except Exception as e:
            raise ValueError(f"Failed to perform Gemini OCR on PDF: {e}")

    def parse_text_via_gemini(self, text: str) -> Dict[str, Any]:
        """Use Gemini 2.5 Flash to parse raw resume text into the required structured JSON schema"""
        if not settings.GEMINI_API_KEY:
            logger.info("Gemini API key is missing. Using regex parser fallback.")
            return self.parse_text_to_schema(text)
            
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={settings.GEMINI_API_KEY}"
        headers = {"Content-Type": "application/json"}
        
        system_prompt = (
            "You are an expert ATS resume parser. Extract structured information from the raw resume text. "
            "Return ONLY a valid JSON object matching this exact schema:\n"
            "{\n"
            "  \"personal_info\": {\n"
            "    \"name\": \"Full Name (look at first few lines, fix spaced letters if any)\",\n"
            "    \"email\": \"email@example.com\",\n"
            "    \"phone\": \"+1234567890\",\n"
            "    \"linkedin\": \"linkedinurl\",\n"
            "    \"portfolio\": \"portfolio or github url\",\n"
            "    \"location\": \"City, Country or Location\"\n"
            "  },\n"
            "  \"education\": [\n"
            "    {\n"
            "      \"title\": \"Degree — Major - Date Range - University | Location\",\n"
            "      \"description\": \"GPA, coursework, honors, or achievements\"\n"
            "    }\n"
            "  ],\n"
            "  \"experience\": [\n"
            "    {\n"
            "      \"title\": \"Job Title — Date Range - Organization | Location\",\n"
            "      \"description\": \"– Bullet point 1\\n– Bullet point 2\"\n"
            "    }\n"
            "  ],\n"
            "  \"projects\": [\n"
            "    {\n"
            "      \"title\": \"Project Name - Date - Context/Organization\",\n"
            "      \"description\": \"– Bullet describing project impact\"\n"
            "    }\n"
            "  ],\n"
            "  \"certifications\": [\n"
            "    \"Certification Name — Issuer Date\"\n"
            "  ],\n"
            "  \"skills\": [\n"
            "    \"Skill 1\", \"Skill 2\", \"Skill 3\"\n"
            "  ]\n"
            "}\n"
            "CRITICAL Rules:\n"
            "1. Extract ONLY facts present in the resume text. Do not invent any experience, education, or skills. "
            "If a section is not present in the text, leave it empty.\n"
            "2. Ensure experience description bullet points start with an en-dash (–) or standard bullet.\n"
            "3. Clean up name formatting (e.g., collapse spaces in names like 'A D I T Y A  P U T R A  A F E N D I' to 'Aditya Putra Afendi').\n"
            "4. Return ONLY JSON. No markdown backticks or wrappers."
        )
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": system_prompt},
                        {"text": f"Raw Resume Text:\n{text}"}
                    ]
                }
            ],
            "generationConfig": {
                "responseMimeType": "application/json",
                "temperature": 0.1
            }
        }
        
        try:
            with httpx.Client() as client:
                r = client.post(url, json=payload, headers=headers, timeout=25.0)
                if r.status_code == 200:
                    data = r.json()
                    response_text = data["contents"][0]["parts"][0]["text"].strip()
                    if response_text.startswith("```json"):
                        response_text = response_text[7:]
                    if response_text.endswith("```"):
                        response_text = response_text[:-3]
                    return json.loads(response_text.strip())
                else:
                    logger.error(f"Gemini Parser API returned code {r.status_code}: {r.text}")
        except Exception as e:
            logger.error(f"Gemini Parser API invocation error: {e}. Falling back to regex parser.")
            
        return self.parse_text_to_schema(text)

    def parse_resume(self, file_bytes: bytes, file_type: str) -> Dict[str, Any]:
        """Parse resume file bytes and return structured data including raw_text."""
        if "pdf" in file_type.lower():
            text = self.extract_text_from_pdf(file_bytes)
            if len(text.strip()) < 150:
                logger.info("PDF has very little selectable text. Running Gemini OCR.")
                try:
                    text = self.extract_text_via_gemini_ocr(file_bytes)
                except Exception as e:
                    logger.error(f"Gemini OCR fallback failed: {e}. Proceeding with raw extracted text.")
        elif "docx" in file_type.lower() or "officedocument" in file_type.lower():
            text = self.extract_text_from_docx(file_bytes)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are accepted.")

        structured = self.parse_text_via_gemini(text)
        # Store the clean raw text so it can be sent directly to Claude Sonnet
        # without re-parsing artifacts. This is the "ground truth" CV text.
        structured["raw_text"] = text.strip()
        return structured

    def parse_text_to_schema(self, text: str) -> Dict[str, Any]:
        """Convert raw resume text into structured schema via section slicing"""
        schema = {
            "personal_info": {
                "name": "",
                "email": "",
                "phone": "",
                "linkedin": "",
                "portfolio": "",
                "location": ""
            },
            "education": [],
            "experience": [],
            "projects": [],
            "certifications": [],
            "skills": []
        }

        # 1. Parse Personal Info (using regex)
        # Find Email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            schema["personal_info"]["email"] = email_match.group(0).strip()

        # Find Phone
        phone_match = re.search(r'(\+?\d[\d\-\s\(\)\.]{7,16}\d)', text)
        if phone_match:
            schema["personal_info"]["phone"] = phone_match.group(0).strip()

        # Find LinkedIn
        linkedin_match = re.search(r'(linkedin\.com/in/[\w\-]+)', text, re.IGNORECASE)
        if linkedin_match:
            schema["personal_info"]["linkedin"] = linkedin_match.group(0).strip()

        # Find GitHub / Portfolio
        portfolio_match = re.search(r'((github\.com/[\w\-]+)|(portfolio\.[\w\.]+)|([\w\-]+\.me)|([\w\-]+\.github\.io))', text, re.IGNORECASE)
        if portfolio_match:
            schema["personal_info"]["portfolio"] = portfolio_match.group(0).strip()

        # Try to extract name (usually first 1-2 lines)
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        if lines:
            # Simple heuristic: first line is often the name
            # Make sure it's not an email or address
            candidate = lines[0]
            if "@" not in candidate and len(candidate) < 80:
                # Collapse spaced-out names like "A D I T Y A  P U T R A  A F E N D I"
                schema["personal_info"]["name"] = self._collapse_spaced_name(candidate)

        # 2. Slice sections based on standard header matches
        sections = {
            "education": ["education", "academic history", "academic background", "study", "studies", "kuliah"],
            "experience": ["experience", "employment history", "work history", "professional experience", "career", "pengalaman kerja", "work experience", "additional leadership & organizational experience", "additional leadership and organizational experience", "leadership experience", "leadership"],
            "projects": ["projects", "personal projects", "key projects", "proyek", "notable projects"],
            "certifications": ["certifications", "certificates", "licenses", "sertifikasi", "sertifikat", "key certifications", "notable achievements", "achievements"],
            "skills": ["skills", "technical skills", "core competencies", "expertise", "keahlian", "skills & interests", "skills and interests"]
        }

        current_section = None
        section_texts = {sec: [] for sec in sections}

        for line in lines:
            lower_line = line.lower().strip()
            # Normalize by removing all spaces and special characters
            normalized_line = re.sub(r'[^a-z0-9]', '', lower_line)
            
            # Check if this line is a heading
            found_heading = False
            for sec, keywords in sections.items():
                # Normalize keywords too
                normalized_keywords = [re.sub(r'[^a-z0-9]', '', kw) for kw in keywords]
                if normalized_line in normalized_keywords:
                    current_section = sec
                    found_heading = True
                    break
            
            if found_heading:
                continue
                
            if current_section:
                section_texts[current_section].append(line)

        # 3. Format sliced sections
        # Education parsing
        if section_texts["education"]:
            schema["education"] = self._parse_items_by_bullet(section_texts["education"])
            
        # Experience parsing
        if section_texts["experience"]:
            schema["experience"] = self._parse_items_by_bullet(section_texts["experience"])
            
        # Projects parsing
        if section_texts["projects"]:
            schema["projects"] = self._parse_items_by_bullet(section_texts["projects"])

        # Certifications parsing — filter out bare dates and achievement-like items
        if section_texts["certifications"]:
            for line in section_texts["certifications"]:
                clean = line.strip()
                if clean:
                    schema["certifications"].append(clean)

        # Skills parsing — handle "Category: skill · skill · skill" format
        if section_texts["skills"]:
            skills_set = set()
            for line in section_texts["skills"]:
                # Remove bold category prefix like "Product Management: "
                # The category itself is useful, but the individual skills matter more
                # Split by common separators: · , | • - *
                splits = re.split(r'[·,|•\*]', line)
                for s in splits:
                    cleaned = s.strip()
                    # Remove leading category label if present ("Category: value")
                    if ':' in cleaned:
                        parts = cleaned.split(':', 1)
                        category = parts[0].strip()
                        value = parts[1].strip()
                        # Add both category and value if meaningful
                        if category and len(category) < 40 and len(category) > 2:
                            skills_set.add(category)
                        if value and len(value) < 40 and len(value) > 2:
                            skills_set.add(value)
                    elif cleaned and len(cleaned) < 40 and len(cleaned) > 2:
                        # Filter out noise words
                        if cleaned.lower() not in ('to', 'and', 'or', 'the', 'a', 'an', 'of', 'in', 'end'):
                            skills_set.add(cleaned)
            schema["skills"] = sorted(list(skills_set))

        return schema

    def _parse_items_by_bullet(self, lines: List[str]) -> List[Dict[str, Any]]:
        """Group text lists into objects (e.g. mapping bullets to titles and descriptions)"""
        items = []
        current_item = {"title": "", "description": []}
        header_lines_count = 0
        has_bullets = False
        
        n = len(lines)
        for idx in range(n):
            line = lines[idx]
            is_bullet = line.startswith("•") or line.startswith("-") or line.startswith("*") or line.startswith("▪") or line.startswith("–")
            clean_line = re.sub(r'^[•\-\*▪–]\s*', '', line).strip()
            
            if is_bullet:
                if current_item["title"]:
                    current_item["description"].append(clean_line)
                    has_bullets = True
                else:
                    current_item["title"] = clean_line
                    header_lines_count = 1
                    has_bullets = True
            else:
                # Non-bullet line
                if not current_item["title"]:
                    current_item["title"] = clean_line
                    header_lines_count = 1
                else:
                    # If we already have a description block, this must be a new item
                    if current_item["description"]:
                        # Determine if this starts a new item or is a continuation of the last bullet
                        is_capitalized = clean_line[0].isupper() if clean_line else False
                        has_date = bool(re.search(r'(20\d\d|19\d\d|present|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|current)', clean_line, re.IGNORECASE))
                        is_separator = any(sep in clean_line for sep in ["|", "@", "—", " - ", " – "]) or (" at " in f" {clean_line.lower()} ")
                        
                        # Lookahead to see if next line has a date (common for titles followed by date line)
                        next_has_date = False
                        if idx + 1 < n:
                            next_line = lines[idx + 1]
                            next_has_date = bool(re.search(r'(20\d\d|19\d\d|present|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|current)', next_line, re.IGNORECASE))
                            
                        if is_capitalized and (has_date or is_separator or next_has_date):
                            # Save current item
                            items.append({
                                "title": current_item["title"],
                                "description": "\n".join(current_item["description"])
                            })
                            current_item = {"title": clean_line, "description": []}
                            header_lines_count = 1
                            has_bullets = False
                        else:
                            # Continuation of the last bullet
                            if current_item["description"]:
                                current_item["description"][-1] += " " + clean_line
                            else:
                                current_item["description"].append(clean_line)
                    else:
                        # If we have less than 3 header lines, and it looks like a header (short or has date), append to title
                        has_date = bool(re.search(r'(20\d\d|19\d\d|present|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|current)', clean_line, re.IGNORECASE))
                        is_header_like = (len(clean_line) < 60) or has_date or ("|" in clean_line)
                        
                        if header_lines_count < 3 and is_header_like:
                            current_item["title"] += " - " + clean_line
                            header_lines_count += 1
                        else:
                            # Put it in description
                            current_item["description"].append(clean_line)
                            
        if current_item["title"]:
            items.append({
                "title": current_item["title"],
                "description": "\n".join(current_item["description"])
            })
            
        return items

    def _collapse_spaced_name(self, name: str) -> str:
        """Collapse spaced-out names like 'A D I T Y A  P U T R A  A F E N D I' to 'Aditya Putra Afendi'.
        
        The input uses single spaces between letters within a word, and double spaces between words.
        """
        if not name:
            return name
        
        # Check if the name looks spaced-out: mostly single characters separated by spaces
        # A quick test: count single-letter tokens vs multi-letter tokens
        tokens = name.split()
        if len(tokens) > 4:
            single_chars = sum(1 for t in tokens if len(t) == 1)
            if single_chars > len(tokens) * 0.5:
                # Split by double-space (or more) to get word groups
                # Each word group is like "A D I T Y A"
                word_groups = re.split(r'\s{2,}', name.strip())
                collapsed_words = []
                for group in word_groups:
                    # Each group is single letters separated by single spaces: "A D I T Y A"
                    letters = group.replace(' ', '')
                    if letters:
                        collapsed_words.append(letters.title())
                if collapsed_words:
                    return ' '.join(collapsed_words)
        
        return name

resume_parser = ResumeParser()

